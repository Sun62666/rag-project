"""文档上传与向量化 API"""
import os
import logging
import hashlib
import shutil
from typing import List
from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from pydantic import BaseModel
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.embeddings import DashScopeEmbeddings

from src.api.deps import get_current_user_dep
from src.core.config import get_settings

logger = logging.getLogger(__name__)

router = APIRouter(tags=["运维工具-文档"])

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
UPLOAD_DIR = os.path.join(PROJECT_ROOT, "data", "uploads")


def _normalize_uri(uri: str) -> str:
    if not uri.startswith("http://") and not uri.startswith("https://"):
        uri = f"http://{uri}"
    return uri


# ==================== 数据模型 ====================

class UploadResponse(BaseModel):
    filename: str
    chunks: int
    status: str
    message: str = ""


# ==================== 接口 ====================

@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    collection_name: str = Form(default=""),
    _user: str = Depends(get_current_user_dep),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".pdf", ".txt", ".md", ".docx"]:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}, 仅支持 pdf/txt/md/docx")

    os.makedirs(UPLOAD_DIR, exist_ok=True)
    save_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(save_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        docs = _load_file(save_path)
        if not docs:
            os.remove(save_path)
            return UploadResponse(filename=file.filename, chunks=0, status="error", message="文件内容为空或无法解析")

        cfg = get_settings()
        splits = _split_and_inject(docs)
        col_name = collection_name or cfg.COLLECTION_NAME

        _append_to_milvus(splits, col_name, cfg)

        # 如果上传到通用文档库，重建 BM25 索引
        if col_name == "property_regulations":
            from src.tools.document_qa import rebuild_document_qa_bm25
            rebuild_document_qa_bm25(splits)

        return UploadResponse(
            filename=file.filename,
            chunks=len(splits),
            status="ok",
            message=f"成功导入 {len(splits)} 个文档块到集合 {col_name}",
        )
    except Exception as e:
        logger.error(f"文档上传处理失败: {e}")
        return UploadResponse(filename=file.filename, chunks=0, status="error", message=str(e))


@router.get("/upload/list")
async def list_uploaded_docs(_user: str = Depends(get_current_user_dep)):
    if not os.path.isdir(UPLOAD_DIR):
        return {"files": []}
    files = []
    for f in os.listdir(UPLOAD_DIR):
        fp = os.path.join(UPLOAD_DIR, f)
        if os.path.isfile(fp):
            ext = os.path.splitext(f)[1].lower()
            if ext in [".pdf", ".txt", ".md", ".docx"]:
                files.append({
                    "name": f,
                    "size": os.path.getsize(fp),
                    "modified": os.path.getmtime(fp),
                    "type": ext[1:],
                })
    return {"files": files}


@router.delete("/upload/{filename}")
async def delete_uploaded_doc(filename: str, _user: str = Depends(get_current_user_dep)):
    fp = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(fp):
        raise HTTPException(status_code=404, detail="文件不存在")
    os.remove(fp)
    return {"status": "ok", "message": f"已删除 {filename}"}


@router.get("/knowledge/stats")
async def knowledge_stats(_user: str = Depends(get_current_user_dep)):
    cfg = get_settings()
    uri = _normalize_uri(cfg.MILVUS_URI)
    try:
        from pymilvus import MilvusClient
        client = MilvusClient(uri=uri)
        collections = client.list_collections()
        result = {}
        for col in collections:
            if client.has_collection(col):
                stats = client.get_collection_stats(col)
                result[col] = int(stats.get("row_count", 0))
        client.close()
        return {"collections": result}
    except Exception as e:
        return {"error": str(e), "collections": {}}


# ==================== 内部函数 ====================

def _load_file(path: str) -> List[Document]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return PyPDFLoader(path).load()
    elif ext in [".txt", ".md"]:
        return TextLoader(path, encoding="utf-8").load()
    elif ext == ".docx":
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            return Docx2txtLoader(path).load()
        except ImportError:
            # fallback: 使用 python-docx
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            return [Document(page_content=text, metadata={"source": os.path.basename(path)})]
    return []


def _split_and_inject(docs: List[Document], chunk_size: int = 500, chunk_overlap: int = 60) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n案例 ", "\n案例", "案例", "\n问：", "\n问:", "\n## ", "\n\n", "\n", "。", " ", ""],
    )
    splits = splitter.split_documents(docs)
    for doc in splits:
        doc_id = hashlib.md5(doc.page_content.encode()).hexdigest()[:12]
        doc.metadata.setdefault("doc_id", doc_id)
    return splits


def _append_to_milvus(splits: List[Document], collection_name: str, cfg):
    from src.core.milvus_compat import ensure_milvus_connection
    uri = _normalize_uri(cfg.MILVUS_URI)
    emb = DashScopeEmbeddings(model=cfg.EMBED_MODEL, dashscope_api_key=cfg.DASHSCOPE_API_KEY)

    client = ensure_milvus_connection(uri)
    exists = client.has_collection(collection_name)

    from langchain_milvus import Milvus as MilvusVS

    if exists:
        vs = MilvusVS(
            embedding_function=emb,
            collection_name=collection_name,
            connection_args={"uri": uri},
            auto_id=True,
            enable_dynamic_field=True
        )
        vs.add_documents(splits)
    else:
        MilvusVS.from_documents(
            splits,
            emb,
            collection_name=collection_name,
            connection_args={"uri": uri},
            auto_id=True,
            enable_dynamic_field=True
        )
